"""
migrate_resting_to_groundstate.py
=================================
Finish the "resting" -> "ground-state" wording migration that was started
2026-06-16/17 but never completed (and that today's inserts regressed). Every
"resting X" adjective phrase becomes "ground-state X", matching the manuscript's
existing hyphenated "ground-state geometry / ground-state chromophore" style.

Run-level text replacement across all body paragraphs AND table cells (Table 4
and the Table 1 caption both contain "resting position"). Backs up first.
Prose only -> verify_manuscript_numbers.py unaffected.
"""
from __future__ import annotations
import shutil
from pathlib import Path

import docx

PROJECT = Path(__file__).resolve().parent.parent
SRC = PROJECT / "manuscript-complete-mz.docx"
BACKUP = PROJECT / "manuscript-complete-mz_backup-groundstate.docx"

# ordered: hyphenated compound first, then capitalized, then bare lowercase
REPLACEMENTS = [
    ("resting-position", "ground-state position"),  # drop the double hyphen
    ("Resting", "Ground-state"),                      # block-B heading "Resting twist"
    ("resting", "ground-state"),
]


def replace_in_runs(paragraph) -> int:
    n = 0
    for r in paragraph.runs:
        t = r.text
        if "resting" in t.lower():
            new = t
            for old, rep in REPLACEMENTS:
                new = new.replace(old, rep)
            if new != t:
                r.text = new
                n += 1
    return n


def main():
    shutil.copy(SRC, BACKUP)
    d = docx.Document(str(SRC))
    n = 0
    for p in d.paragraphs:
        n += replace_in_runs(p)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += replace_in_runs(p)
    d.save(str(SRC))

    # report leftovers
    txt = " ".join(p.text for p in docx.Document(str(SRC)).paragraphs)
    for t in docx.Document(str(SRC)).tables:
        for row in t.rows:
            for cell in row.cells:
                txt += " " + cell.text
    left = txt.lower().count("resting")
    gs = txt.lower().count("ground-state") + txt.lower().count("ground state")
    print(f"runs changed: {n}; remaining 'resting': {left}; 'ground-state' now: {gs}")
    print(f"backup: {BACKUP.name}")


if __name__ == "__main__":
    main()
