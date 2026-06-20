#!/usr/bin/env python3
"""Reorder the Abstract and Introduction of manuscript-complete-mz.docx so the
findings appear in the same sequence and emphasis as Results & Discussion:
(A) every barrel clamps the I-bond / the P-bond is the variable degree of
freedom; (B) position 203 is the dominant P-bond gatekeeper; (C) ground-state
twist (resting position) predicts brightness in red FPs; (D) cage size reports
chromophore chemistry, not brightness. Edits four plain-run paragraphs in place.
Run from the project root.
"""
import docx

DOC = "manuscript-complete-mz.docx"

# (current-text prefix used to locate the paragraph, new full text)
EDITS = [
    # Abstract paragraph 2: scan setup + (A) + (B)
    ("We scanned 838 structures, measuring two geometric",
     "We scanned 838 structures, measuring two geometric properties of each "
     "chromophore environment: the fraction of the (τ, φ) torsional "
     "sphere that is sterically accessible to the chromophore, and the angular "
     "distance from the deposited chromophore to the nearest planar geometry. "
     "Every barrel imposes the same asymmetry on the two single bonds that flank "
     "the methine bridge: the I-bond, the excited-state cis-trans isomerization "
     "axis, is clamped to roughly two of seventy-two orientations in all color "
     "classes, whereas the P-bond, the phenol flip, is the variable degree of "
     "freedom. Sweeping each chromophore's torsions until first collision "
     "identifies position 203 (avGFP numbering) as the dominant P-bond "
     "gatekeeper, the first steric barrier in 21% of all sweep events and "
     "2.5-fold more often than any other position; the T203Y substitution that "
     "gave rise to the YFP/Citrine lineage is a historical example consistent "
     "with this gatekeeper model."),

    # Abstract paragraph 3: (C) + (D) + engineering message
    ("The engineering message is that FP design",
     "What this restraint controls, however, is not rotational room but resting "
     "geometry. Among red FPs a more twisted ground-state chromophore is "
     "consistently dimmer, whereas cage size itself tracks chromophore chemistry "
     "rather than brightness: indole-based cyan and acylimine-based red FPs have "
     "the tightest cages and imidazole-based blue FPs the loosest, yet within a "
     "color class cage size does not predict quantum yield. These two quantities "
     "are nearly uncorrelated and diverge as predictors of brightness. The "
     "engineering message is that FP design should focus less on making the cage "
     "generally tight or roomy and more on where the protein positions the "
     "chromophore."),

    # Introduction opening: thesis reordered to lead with the asymmetry (A), then
    # the brightness predictor (C), then the cage-chemistry result (D)
    ("The brightness, or quantum yield (QY), of a fluorescent",
     "The brightness, or quantum yield (QY), of a fluorescent protein is widely "
     "believed to track how rigidly the barrel confines the chromophore "
     "[1, 2, 3, 4, 5]. That belief has been examined one protein at a time "
     "[4, 5, 6, 7] and probed by recent simulations [8, 9, 10, 11], but never "
     "systematically across the entire structural family. This paper tests it on "
     "838 crystal structures spanning the full range of chromophore chemistries "
     "and color classes in the Protein Databank, using a purely geometric rigid "
     "scan. The result is that every barrel clamps the isomerization-driving "
     "I-bond while leaving the P-bond as the variable element, that the deposited "
     "chromophore's distance from planarity predicts quantum yield among red FPs, "
     "and that cage size reports chromophore chemistry rather than within-class "
     "brightness. This changes the fluorescent protein engineering message."),

    # Introduction "what we find": full A -> B -> C -> D
    ("Here we show that cage size is a strong reporter",
     "Here we show that every fluorescent protein barrel imposes the same "
     "asymmetry on the two torsions flanking the methine bridge: the I-bond is "
     "clamped in all color classes, while the P-bond is the variable degree of "
     "freedom. Our gatekeeper analysis identifies position 203 as the dominant "
     "local constraint on P-bond rotation, making this long-recognized "
     "color-tuning site a promising target for engineering chromophore behavior. "
     "Crystallographic chromophore geometry, rather than the amount of accessible "
     "space, predicts quantum yield among red FPs, where a more twisted "
     "ground-state chromophore is consistently dimmer. Cage size itself proves a "
     "strong reporter of chromophore chemistry but a poor predictor of brightness "
     "within a color class."),
]


def set_text(p, text):
    for r in p.runs[1:]:
        r._element.getparent().remove(r._element)
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def main():
    d = docx.Document(DOC)
    # locate all targets first (by current prefix), then edit
    targets = []
    for prefix, new in EDITS:
        hit = next((p for p in d.paragraphs if p.text.strip().startswith(prefix)), None)
        assert hit is not None, f"paragraph not found: {prefix!r}"
        targets.append((hit, new))
    for p, new in targets:
        set_text(p, new)
    d.save(DOC)
    print(f"Reordered {len(targets)} paragraphs in {DOC}.")


if __name__ == "__main__":
    main()
