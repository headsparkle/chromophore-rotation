"""
bj_reangle.py
=============
Re-angle the manuscript for Biophysical Journal:
  (A) trim the Abstract to <= 250 words (BJ limit), preserving every result;
  (B) add a Significance Statement (<= 50 words, BJ requirement) after the Abstract;
  (C) rename "2. Methods" -> "2. Materials and Methods" (BJ convention);
  (D) upgrade the Methods 2.4 data/code-availability sentence to a proper statement.

Backs up first. Prose only; headline numbers untouched -> verifier unaffected.
"""
from __future__ import annotations
import copy
import shutil
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

PROJECT = Path(__file__).resolve().parent.parent
SRC = PROJECT / "manuscript-complete-mz.docx"
BACKUP = PROJECT / "manuscript-complete-mz_backup-bj.docx"

ABSTRACT_P1_OLD = "A common explanation for brightness variation in fluorescent proteins (FPs) is that tight chromophore cages"
ABSTRACT_P1 = (
    "A common explanation for brightness variation in fluorescent proteins (FPs) is that tight chromophore "
    "cages produce bright proteins, predicting that quantum yield should track the dihedral rotational space "
    "available to the chromophore. We tested this across 838 FP crystal structures from the Protein Data Bank."
)
ABSTRACT_P2_OLD = "For each structure, we scanned the chromophore"
ABSTRACT_P2 = (
    "For each structure we scanned the chromophore’s two methine-bridge torsions, τ and φ, measuring "
    "the fraction of torsional space that is sterically accessible and the angular distance from the deposited "
    "chromophore to the nearest planar geometry. The scan reveals a consistent family-wide asymmetry: the "
    "I-bond, the excited-state cis-trans isomerization axis, is clamped in all color classes, while the P-bond, "
    "the phenol-flip axis, remains the variable degree of freedom. Gatekeeper analysis identifies position 203 "
    "(avGFP numbering) as the dominant P-bond constraint, the first steric barrier in 21% of all sweep events."
)
ABSTRACT_P3_OLD = "The key result, however, is that cage size does not generally predict quantum yield."
ABSTRACT_P3 = (
    "Critically, cage size does not generally track quantum yield; instead it reports chromophore chemistry, "
    "with indole-based cyan and acylimine-based red FPs the tightest and imidazole-based blue FPs the loosest. "
    "Within color classes, cage size is a poor predictor of brightness. In red FPs, quantum yield is instead "
    "associated with the chromophore’s ground-state geometry: more twisted crystallographic chromophores "
    "are consistently dimmer. More broadly, the scan shows that rigidity is not a single property: cage size, "
    "ground-state geometry, gatekeeper residues, and thermal immobilization are distinct quantities with "
    "different relationships to quantum yield. Fluorescent-protein design should therefore target local "
    "geometric control of the chromophore rather than global cage tightening."
)

SIGNIFICANCE = (
    "Fluorescent-protein brightness is widely attributed to a rigid chromophore cage. Scanning chromophore "
    "rotational space across 838 crystal structures, we show that rigidity is not one property: cage size "
    "reports chromophore chemistry, whereas ground-state geometry tracks quantum yield in red fluorescent "
    "proteins. This reframes how structure-guided design should target brightness."
)

AVAIL_OLD = "Code and data are archived at [URL]."
AVAIL_NEW = (
    "All analysis code, the PDB list and the non-rotatable exclusion list, the quantum-yield curation table "
    "with provenance, the avGFP numbering map, and the per-structure (SD1) and per-unique-FP (SD2) data tables "
    "are deposited in a public repository archived at Zenodo (DOI to be assigned on acceptance; [URL])."
)


def set_single_run(p: Paragraph, text: str):
    for r in list(p._p.findall(qn("w:r"))):
        p._p.remove(r)
    p.add_run(text)


def main():
    shutil.copy(SRC, BACKUP)
    d = docx.Document(str(SRC))
    paras = d.paragraphs

    def find(prefix):
        for p in paras:
            if p.text.strip().startswith(prefix):
                return p
        raise RuntimeError(f"not found: {prefix!r}")

    # (A) abstract trim
    p1 = find(ABSTRACT_P1_OLD); set_single_run(p1, ABSTRACT_P1)
    p2 = find(ABSTRACT_P2_OLD); set_single_run(p2, ABSTRACT_P2)
    p3 = find(ABSTRACT_P3_OLD); set_single_run(p3, ABSTRACT_P3)
    wc = sum(len(x.split()) for x in (ABSTRACT_P1, ABSTRACT_P2, ABSTRACT_P3))
    assert wc <= 250, f"abstract still {wc} words"
    sig_wc = len(SIGNIFICANCE.split())
    assert sig_wc <= 50, f"significance {sig_wc} words"

    # (B) Significance Statement after abstract P3
    if not any("Significance Statement" in p.text for p in paras):
        abs_head = find("Abstract")  # Heading 2, clone for style
        # heading
        h = copy.deepcopy(abs_head._p)
        for r in list(h.findall(qn("w:r"))):
            h.remove(r)
        p3._p.addnext(h)
        hp = Paragraph(h, p3._parent); hp.add_run("Significance Statement")
        # body paragraph after the heading
        body = copy.deepcopy(p3._p)
        for r in list(body.findall(qn("w:r"))):
            body.remove(r)
        h.addnext(body)
        Paragraph(body, p3._parent).add_run(SIGNIFICANCE)

    # (C) Methods heading rename
    for p in paras:
        if p.text.strip() == "2. Methods":
            set_single_run(p, "2. Materials and Methods")
            break

    # (D) availability sentence
    for p in d.paragraphs:
        for r in p.runs:
            if AVAIL_OLD in r.text:
                r.text = r.text.replace(AVAIL_OLD, AVAIL_NEW)

    d.save(str(SRC))
    print(f"abstract words: {wc} (<=250) | significance words: {sig_wc} (<=50)")
    print(f"backup: {BACKUP.name}")


if __name__ == "__main__":
    main()
