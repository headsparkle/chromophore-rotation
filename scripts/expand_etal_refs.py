#!/usr/bin/env python3
"""
expand_etal_refs.py
===================

Expand the 11 "et al." references in the master docx to full author lists,
retrieved from Crossref by DOI (no author name is invented). Also corrects
three metadata errors found while checking:
  - ref 11 (Pieri):  pages 2170-2182  -> 17646-17658
  - ref 21 (Chen):   article e2421111122 -> e2508094122 (the old number is a
                     different PNAS paper)
  - ref 24 (Manna):  title "...nonradiative decay in red fluorescent proteins"
                     -> the published "Dark-state-mediated photobleaching in
                     mCherry-based red fluorescent proteins"; pages 1124-1132
                     -> 3596-3604 (same paper, mis-cited)

Authors are formatted in Biophysical Journal style: first author inverted
"Family, I. I.,"; later authors "I. I. Family"; last preceded by "and".
The SciPy reference is truncated at its "SciPy 1.0 Contributors" group author
(the named byline), not the full consortium expansion. Backs up first.
"""
from __future__ import annotations

import re
import shutil
import urllib.request
import json
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
MASTER = ROOT / "manuscript-complete-mz-new.docx"
BACKUP = ROOT / "manuscript-complete-mz_backup-expandrefs.docx"

# ref number -> (doi, year, title, journal-abbrev, "vol:pages", stop_after_group)
REFS = {
    2: ("10.1016/s0968-0004(00)89099-4", "1995", "Understanding, improving and using green fluorescent proteins", "Trends Biochem. Sci.", "20:448–455", None),
    9: ("10.1021/acs.jctc.1c00748", "2022", "Prediction of fluorophore brightness in designed mini fluorescence activating proteins", "J. Chem. Theory Comput.", "18:3190–3203", None),
    11: ("10.1021/jacs.4c00458", "2024", "Conical intersection accessibility dictates brightness in red fluorescent proteins", "J. Am. Chem. Soc.", "146:17646–17658", None),
    16: ("10.1038/s41586-020-2649-2", "2020", "Array programming with NumPy", "Nature.", "585:357–362", None),
    17: ("10.1038/s41592-019-0686-2", "2020", "SciPy 1.0: fundamental algorithms for scientific computing in Python", "Nat. Methods.", "17:261–272", "SciPy 1.0 Contributors"),
    20: ("10.1371/journal.pbio.3000936", "2020", "Aequorea's secrets revealed: new fluorescent proteins with unique properties for bioimaging and biosensing", "PLoS Biol.", "18:e3000936", None),
    21: ("10.1073/pnas.2508094122", "2025", "A twisted chromophore powers a turn-on fluorescent protein chloride sensor", "Proc. Natl. Acad. Sci. USA.", "122:e2508094122", None),
    22: ("10.1038/ncomms1738", "2012", "Structure-guided evolution of cyan fluorescent proteins towards a quantum yield of 93%", "Nat. Commun.", "3:751", None),
    23: ("10.1038/s41587-022-01278-2", "2022", "A highly photostable and bright green fluorescent protein", "Nat. Biotechnol.", "40:1132–1142", None),
    24: ("10.1021/acs.jpclett.5c04106", "2026", "Dark-state-mediated photobleaching in mCherry-based red fluorescent proteins", "J. Phys. Chem. Lett.", "17:3596–3604", None),
    27: ("10.1002/pro.3280", "2018", "Improvements to the APBS biomolecular solvation software suite (pdb2pqr)", "Protein Sci.", "27:112–128", None),
}


def norm_hyphen(s: str) -> str:
    # normalize Unicode hyphens (figure/non-breaking) to ASCII so initials and
    # surnames render consistently, e.g. "Guo‐Wei" -> "Guo-Wei"
    return s.replace("‐", "-").replace("‑", "-")


def fetch_authors(doi):
    url = f"https://api.crossref.org/works/{doi}?mailto=mzim@conncoll.edu"
    req = urllib.request.Request(url, headers={"User-Agent": "refcheck/1.0 (mailto:mzim@conncoll.edu)"})
    with urllib.request.urlopen(req, timeout=30) as r:
        return json.load(r)["message"].get("author", [])


def initials(given: str) -> str:
    given = norm_hyphen(given)
    out = []
    for part in (p for p in re.split(r"[\s.]+", given.strip()) if p):
        if "-" in part:
            out.append("-".join(s[0].upper() + "." for s in part.split("-") if s))
        else:
            out.append(part[0].upper() + ".")
    return " ".join(out)


def bj_authors(authors, stop_after_group=None) -> str:
    items = []
    for a in authors:
        if a.get("name"):
            items.append(("GROUP", a["name"]))
            if stop_after_group and a["name"] == stop_after_group:
                break
        else:
            items.append((initials(a.get("given", "")), norm_hyphen(a.get("family", ""))))
    parts = []
    for i, (ini, fam) in enumerate(items):
        if ini == "GROUP":
            parts.append(fam)
        elif i == 0:
            parts.append(f"{fam}, {ini}")
        else:
            parts.append(f"{ini} {fam}")
    if len(parts) == 1:
        return parts[0]
    return ", ".join(parts[:-1]) + ", and " + parts[-1]


def build():
    out = {}
    for n, (doi, yr, title, jour, vp, grp) in REFS.items():
        aj = bj_authors(fetch_authors(doi), grp)
        out[n] = f"{n}. {aj}. {yr}. {title}. {jour} {vp}."
    return out


def set_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)


def main():
    new_refs = build()
    doc = Document(str(MASTER))
    ps = doc.paragraphs
    ref_i = next(i for i, p in enumerate(ps)
                 if p.style and p.style.name == "Heading 2" and p.text.strip() == "References")
    ref_ps = {}
    for p in ps[ref_i + 1:]:
        m = re.match(r"^(\d+)\.\s", p.text)
        if m:
            ref_ps[int(m.group(1))] = p

    shutil.copy2(MASTER, BACKUP)
    print(f"backup -> {BACKUP.name}")
    for n, text in new_refs.items():
        if n not in ref_ps:
            raise SystemExit(f"ref {n} paragraph not found")
        set_text(ref_ps[n], text)
        na = text.count(",") and text.split(". ")[1]
        print(f"  ref {n}: expanded ({text.split('.')[1].count(',') + 1} name-commas) -> {text[:55]}...")
    doc.save(str(MASTER))
    print(f"saved -> {MASTER.name}; expanded {len(new_refs)} references")


if __name__ == "__main__":
    main()
