"""
insert_engineering_table.py
===========================
(1) Correct the two per-class gatekeeper claims that the sweep data contradicts
    (cyan dominant gatekeeper is 203/167, not 146; red does not converge on 203),
    in the gatekeeper paragraph and the Conclusion.
(2) Insert Table 4, a data-faithful prospective engineering guide, into the
    Conclusion after the "practical implication" paragraph.

Backs up first. Prose + one new table; verify_manuscript_numbers.py unaffected.
"""
from __future__ import annotations
import copy
import shutil
from pathlib import Path

import docx
from docx.oxml.ns import qn

PROJECT = Path(__file__).resolve().parent.parent
SRC = PROJECT / "manuscript-complete-mz.docx"
BACKUP = PROJECT / "manuscript-complete-mz_backup-engtable.docx"

# ---- prose corrections (old -> new substring swaps within a run) ----------
SWAPS = [
    # gatekeeper paragraph: scope 203 to green/yellow, add red caveat
    ("The position-203 rule applies to all phenol-containing chromophores (green, yellow, red FPs).",
     "The position-203 rule is established for the phenol-type green and yellow FPs, where Thr203 or Tyr203 is the dominant first-clash wall and is confirmed by the forward deletion test above. Red FPs also carry a phenol ring, but their first-clash gatekeepers do not converge on position 203: the most frequent sweep walls (deposited positions 146, 197, and 163) vary by lineage and are not reliably transferable to avGFP numbering, so red-FP design rests on the resting-position and electrostatic logic developed above rather than on a single gatekeeper site."),
    # gatekeeper paragraph: cyan gatekeepers are 203/167; 146 is the cage lever
    ("For indole-type FPs the dominant gatekeeper is position 146, as we discuss below.",
     "For indole-type cyan FPs the first-clash gatekeepers are instead positions 203 and 167; the cyan brightness lever is not a gatekeeper but cage packing at position 146 (the I146F effect discussed below), which tightens the indole cavity without forming the first rotational wall."),
    # Conclusion: correct the per-class actionable-site sentence
    ("Position 203 is the most actionable site for phenol-containing chromophores, while position 146 plays an analogous role in cyan FPs.",
     "Position 203 is the most actionable site for the green and yellow phenol-type FPs; in cyan FPs the analogous lever is cage packing at position 146, and in red FPs the resting position itself, rather than any single gatekeeper, is the better design handle."),
]

# ---- Table 4 -------------------------------------------------------------
TABLE4_ANCHOR = "The practical implication is that FP design should move from global cage tightening"
TABLE4_CAPTION = (
    "Table 4. Prospective engineering guide by chromophore class. Gatekeeper and "
    "lever sites are given in avGFP numbering. Site assignments are firm for the "
    "phenol-type green and yellow FPs, where position 203 is validated by the "
    "in-silico deletion test; for the red, cyan, blue, and mixed classes the table "
    "reports the data-supported design lever and flags where single-site transfer "
    "is unreliable."
)
TABLE4_HEADER = ["Chromophore class", "Gatekeeper / lever sites (avGFP numbering)",
                 "Design suggestion", "Main caution"]
TABLE4_ROWS = [
    ["Green / yellow (phenol HBI/CRO)",
     "203 validated by in-silico deletion (Thr203/Tyr203); 205 from matched-pair enrichment; 167 secondary",
     "Pack the P-bond wall above the phenol at 203 and 205; in yellows Tyr203 adds a pi-stack",
     "Planarity and local restraint matter more than global cage tightening"],
    ["Red (acylimine / compact NRQ/CRQ)",
     "No convergence on 203; sweep walls (deposited 146, 197, 163) vary by lineage and are not avGFP-transferable",
     "Park the chromophore near planar and stiffen the P-bond; resting position is the design handle",
     "Electrostatics matter at least as much as sterics"],
    ["Cyan (indole SWG)",
     "First-clash gatekeepers 203 and 167; brightness lever is cage packing at 146 (I146F)",
     "Tighten the indole cavity with a bulkier residue at 146",
     "The 203 pi-stack logic of phenol-type FPs does not transfer"],
    ["Blue (imidazole IIC)",
     "None validated",
     "Rely on the family-wide I-bond clamp; treat within-class rules as provisional",
     "Chromophore chemistry differences dominate within-class brightness inference"],
    ["Orange / mixed (CRO/NRQ)",
     "Chemistry-dependent; subclass first",
     "Separate CRO-like from NRQ-like cases before choosing mutations",
     "Mixed-class pooling blurs the design logic"],
]


def main():
    shutil.copy(SRC, BACKUP)
    d = docx.Document(str(SRC))

    # (1) prose swaps
    done = 0
    for old, new in SWAPS:
        hit = False
        for p in d.paragraphs:
            if old not in p.text:
                continue
            for r in p.runs:
                if old in r.text:
                    r.text = r.text.replace(old, new)
                    hit = True
                    done += 1
                    break
            if hit:
                break
        if not hit:
            raise RuntimeError(f"prose anchor not found in a single run: {old[:60]!r}")
    print(f"prose swaps applied: {done}/{len(SWAPS)}")

    # (2) Table 4 after the practical-implication paragraph
    anchor = next(p for p in d.paragraphs if TABLE4_ANCHOR in p.text)
    if not any("Table 4." in p.text for p in d.paragraphs):
        # caption paragraph (cloned from anchor for style), bold "Table 4." lead
        cap_p = copy.deepcopy(anchor._p)
        for r in list(cap_p.findall(qn("w:r"))):
            cap_p.remove(r)
        anchor._p.addnext(cap_p)
        from docx.text.paragraph import Paragraph
        cap = Paragraph(cap_p, anchor._parent)
        rb = cap.add_run("Table 4. ")
        rb.bold = True
        cap.add_run(TABLE4_CAPTION[len("Table 4. "):])

        tbl = d.add_table(rows=1, cols=len(TABLE4_HEADER))
        tbl.style = "Table Grid"
        for c, h in zip(tbl.rows[0].cells, TABLE4_HEADER):
            c.text = h
            for para in c.paragraphs:
                for run in para.runs:
                    run.bold = True
        for row in TABLE4_ROWS:
            cells = tbl.add_row().cells
            for c, val in zip(cells, row):
                c.text = val
        # move the table to sit right after the caption
        cap._p.addnext(tbl._tbl)
        print("Table 4 inserted")
    else:
        print("Table 4 already present, skipped")

    d.save(str(SRC))
    print(f"backup: {BACKUP.name}")


if __name__ == "__main__":
    main()
