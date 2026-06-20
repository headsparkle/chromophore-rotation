#!/usr/bin/env python3
"""
convert_intext_citations_bj.py
==============================

Convert in-text citations in the master docx from ACS bracket style [N] to
Biophysical Journal parenthetical style (N). Handles single [N], multiple
[N,M] / [N, M], and ranges [N-M] (range hyphen -> en-dash, e.g. [8-11] ->
(8–11)).

Only purely-numeric citation brackets with 1-2 digit numbers are converted.
This deliberately leaves alone:
  - [URL]                          (Zenodo placeholder)
  - [-0.41, -0.30] etc.            (confidence intervals: minus signs/decimals)
  - [-180, +180)                   (interval notation)
because the regex requires a digit immediately inside the bracket and at most
two digits per number, which none of those satisfy.

Replacement is run-level, so paragraph/run formatting (italics, etc.) is
preserved. Citations were verified to sit within single runs. Table cells are
processed too. A backup is written first, and the script asserts that no
numeric citation bracket remains in the body afterward.
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
MASTER = ROOT / "manuscript-complete-mz-new.docx"
BACKUP = ROOT / "manuscript-complete-mz_backup-intextcite.docx"

CITE = re.compile(r"\[\s*(\d{1,2}(?:\s*[,–-]\s*\d{1,2})*)\s*\]")


def repl(m: re.Match) -> str:
    inner = m.group(1).replace("-", "–")  # range hyphen -> en-dash
    return "(" + inner + ")"


def convert_runs(paragraphs, count):
    for p in paragraphs:
        for r in p.runs:
            if "[" in r.text:
                new = CITE.sub(repl, r.text)
                if new != r.text:
                    count[0] += len(CITE.findall(r.text))
                    r.text = new
    return count


def main():
    doc = Document(str(MASTER))
    ps = doc.paragraphs
    ref_i = next(i for i, p in enumerate(ps)
                 if p.style and p.style.name == "Heading 2"
                 and p.text.strip() == "References")
    body = ps[:ref_i]

    shutil.copy2(MASTER, BACKUP)
    print(f"backup -> {BACKUP.name}")

    count = [0]
    convert_runs(body, count)
    # table cells
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                convert_runs(cell.paragraphs, count)
    print(f"converted {count[0]} in-text citations [N] -> (N)")

    doc.save(str(MASTER))

    # verify nothing numeric-bracketed remains in the body or tables
    doc2 = Document(str(MASTER))
    ps2 = doc2.paragraphs
    ref_i2 = next(i for i, p in enumerate(ps2)
                  if p.style and p.style.name == "Heading 2"
                  and p.text.strip() == "References")
    leftover = []
    for p in ps2[:ref_i2]:
        leftover += CITE.findall(p.text)
    for t in doc2.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    leftover += CITE.findall(p.text)
    if leftover:
        print(f"WARNING: {len(leftover)} citation bracket(s) still present: {leftover}")
    else:
        print("verified: no numeric citation bracket remains in body/tables")


if __name__ == "__main__":
    main()
