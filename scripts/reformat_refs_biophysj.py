#!/usr/bin/env python3
"""
reformat_refs_biophysj.py
=========================

Reformat the manuscript reference list from ACS style to Biophysical Journal
style in the master docx. BJ style:

    N. Lastname, F. M., G. H. Author, and I. J. Author. Year. Sentence-case
       title. Journal Abbrev. Volume:pages.

  - number followed by a period (not [N])
  - first author inverted "Lastname, F. M.,"; later authors "F. M. Lastname";
    last author preceded by "and"; "et al." kept where the source truncated
  - year after the author list
  - journal abbreviation followed by a period, then Volume:pages (en-dash)

Each reference is one Normal paragraph after the References heading. The new
strings below were built by hand from the existing entries so edge cases
(preprint, article numbers, conference proceedings, et al.) are correct.
Backs up first, then replaces paragraph text in place (style preserved).
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
MASTER = ROOT / "manuscript-complete-mz-new.docx"
BACKUP = ROOT / "manuscript-complete-mz_backup-bjrefs.docx"

BJ_REFS = [
    "1. Tsien, R. Y. 1998. The green fluorescent protein. Annu. Rev. Biochem. 67:509–544.",
    "2. Cubitt, A. B., et al. 1995. Understanding, improving and using green fluorescent proteins. Trends Biochem. Sci. 20:448–455.",
    "3. Meech, S. R. 2009. Excited state reactions in fluorescent proteins. Chem. Soc. Rev. 38:2922–2934.",
    "4. Megley, C. M., L. A. Dickson, S. L. Maddalo, G. J. Chandler, and M. Zimmer. 2009. Photophysics and dihedral freedom of the chromophore in yellow, blue, and green fluorescent protein. J. Phys. Chem. B. 113:302–308.",
    "5. Maddalo, S. L., and M. Zimmer. 2006. The role of the protein matrix in green fluorescent protein fluorescence. Photochem. Photobiol. 82:367–372.",
    "6. Reuter, N., H. Lin, and W. Thiel. 2002. Green fluorescent proteins: empirical force field for the neutral and deprotonated forms of the chromophore. J. Phys. Chem. B. 106:6310–6321.",
    "7. Olsen, S., and S. C. Smith. 2008. Bond selection in the photoisomerization reaction of anionic green fluorescent protein and kindling fluorescent protein chromophore models. J. Am. Chem. Soc. 130:8677–8689.",
    "8. Park, J. W., and Y. M. Rhee. 2016. Electric field keeps chromophore planar and produces high yield fluorescence in green fluorescent protein. J. Am. Chem. Soc. 138:13619–13629.",
    "9. Hostetter, E. R., et al. 2022. Prediction of fluorophore brightness in designed mini fluorescence activating proteins. J. Chem. Theory Comput. 18:3190–3203.",
    "10. Jones, C. M., N. H. List, and T. J. Martinez. 2022. Steric and electronic origins of fluorescence in GFP and GFP-like proteins. J. Am. Chem. Soc. 144:12732–12746.",
    "11. Pieri, E., et al. 2024. Conical intersection accessibility dictates brightness in red fluorescent proteins. J. Am. Chem. Soc. 146:2170–2182.",
    "12. Begg, L. P., M. L. Mason, and M. Zimmer. 2026. Barrel shape and chromophore rigidity predict fluorescent-protein photophysics. ChemRxiv https://doi.org/10.26434/chemrxiv.15003310/v2.",
    "13. Wojdyr, M. 2022. GEMMI: a library for structural biology. J. Open Source Softw. 7:4200.",
    "14. Bondi, A. 1964. van der Waals volumes and radii. J. Phys. Chem. 68:441–451.",
    "15. Word, J. M., S. C. Lovell, T. H. LaBean, H. C. Taylor, M. E. Zalis, B. K. Presley, J. S. Richardson, and D. C. Richardson. 1999. Visualizing and quantifying molecular goodness-of-fit: small-probe contact dots with explicit hydrogens. J. Mol. Biol. 285:1711–1733.",
    "16. Harris, C. R., et al. 2020. Array programming with NumPy. Nature. 585:357–362.",
    "17. Virtanen, P., et al. 2020. SciPy 1.0: fundamental algorithms for scientific computing in Python. Nat. Methods. 17:261–272.",
    "18. Seabold, S., and J. Perktold. 2010. statsmodels: econometric and statistical modeling with Python. Proc. 9th Python Sci. Conf.",
    "19. Chen, C., H. Zhang, J. Zhang, H. Ai, and C. Fang. 2023. Structural origin and rational development of bright red noncanonical variants of green fluorescent protein. Phys. Chem. Chem. Phys. 25:15624–15634.",
    "20. Lambert, G. G., et al. 2020. Aequorea's secrets revealed: new fluorescent proteins with unique properties for bioimaging and biosensing. PLoS Biol. 18:e3000936.",
    "21. Chen, X., et al. 2025. A twisted chromophore powers a turn-on fluorescent protein chloride sensor. Proc. Natl. Acad. Sci. USA. 122:e2421111122.",
    "22. Goedhart, J., et al. 2012. Structure-guided evolution of cyan fluorescent proteins towards a quantum yield of 93%. Nat. Commun. 3:751.",
    "23. Hirano, M., et al. 2022. A highly photostable and bright green fluorescent protein. Nat. Biotechnol. 40:1132–1142.",
    "24. Manna, P., et al. 2026. Dark-state-mediated nonradiative decay in red fluorescent proteins. J. Phys. Chem. Lett. 17:1124–1132.",
    "25. Baffour-Awuah, N. Y. A., and M. Zimmer. 2004. Hula-twisting in green fluorescent protein. Chem. Phys. 303:7–11.",
    "26. Ong, W. J.-H., S. Alvarez, I. E. Leroux, R. S. Shahid, A. A. Samma, P. Peshkepija, A. L. Morgan, S. Mulcahy, and M. Zimmer. 2011. Function and structure of GFP-like proteins in the protein data bank. Mol. BioSyst. 7:984–992.",
    "27. Jurrus, E., et al. 2018. Improvements to the APBS biomolecular solvation software suite (pdb2pqr). Protein Sci. 27:112–128.",
    "28. Olsson, M. H. M., C. R. Sondergaard, M. Rostkowski, and J. H. Jensen. 2011. PROPKA3: consistent treatment of internal and surface residues in empirical pKa predictions. J. Chem. Theory Comput. 7:525–537.",
    "29. Heim, R., and R. Y. Tsien. 1996. Engineering green fluorescent protein for improved brightness, longer wavelengths and fluorescence resonance energy transfer. Curr. Biol. 6:178–182.",
]


def set_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)


def main():
    doc = Document(str(MASTER))
    ps = doc.paragraphs
    # locate reference paragraphs: non-empty paragraphs after References heading
    start = None
    for i, p in enumerate(ps):
        if p.style and p.style.name == "Heading 2" and p.text.strip() == "References":
            start = i
            break
    if start is None:
        raise SystemExit("References heading not found")
    ref_ps = [p for p in ps[start + 1:] if p.text.strip()]
    if len(ref_ps) != len(BJ_REFS):
        raise SystemExit(f"expected {len(BJ_REFS)} refs, found {len(ref_ps)}")
    # sanity: confirm order matches [1]..[29]
    for k, p in enumerate(ref_ps, 1):
        if not p.text.strip().startswith(f"[{k}]"):
            raise SystemExit(f"ref {k} mismatch: {p.text[:40]!r}")

    shutil.copy2(MASTER, BACKUP)
    print(f"backup -> {BACKUP.name}")
    for p, new in zip(ref_ps, BJ_REFS):
        set_text(p, new)
    doc.save(str(MASTER))
    print(f"reformatted {len(BJ_REFS)} references to Biophysical Journal style")


if __name__ == "__main__":
    main()
