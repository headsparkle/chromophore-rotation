#!/usr/bin/env python3
"""Build manuscript-complete.docx = manuscript-mz.docx front (through Table 2/
color table, with fixes) + manuscript.docx tail (cyan para -> References),
reconciled to the 3-figure scheme and corrected figure/table/reference order.

Run from the project root.
"""
import copy
import re
import shutil
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

# Figure S3 (Section S1) reference sentence and caption.
FIGS3_REF = (" We reproduce and extend this deposited-torsion map for the full "
             "838-structure cohort in Figure S3.")
FIGS3_CAPTION = [
    ("Figure S3. Deposited chromophore torsions across the fluorescent-protein "
     "family.", True),
    (" Megley torsions φ (P-bond, CA2-CB2-CG2-CD1) versus τ (I-bond, "
     "N2-CA2-CB2-CG2) for all 838 crystal structures, with the phenol "
     "180-degree symmetry folded into φ within (-90, 90] degrees. Points "
     "are colored by color class; gray points lack a curated color assignment. "
     "The deposited geometries cluster near planar (τ ≈ φ ≈ "
     "0) and follow the negative hula-twist diagonal of slope -0.85 reported by "
     "Ong et al. [26] (dashed line). avGFP (1EMA, black star) sits inside the "
     "near-planar core; AausFP2 (6S68, magenta diamond), the dimmest and most "
     "twisted case (QY = 0, 45.7 degrees from planar), lies far outside it.",
     False),
]

SRC_FRONT = "manuscript-mz.docx"     # master front
SRC_TAIL = "manuscript.docx"         # source for everything after the cage section
REF_SRC = "manuscript_pre-merge_backup.md"  # known-correct reference text
OUT = "manuscript-complete.docx"

_TOK = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def parse_inline(s):
    """Split markdown with **bold** / *italic* into (text, bold, italic) runs."""
    parts, pos = [], 0
    for m in _TOK.finditer(s):
        if m.start() > pos:
            parts.append((s[pos:m.start()], False, False))
        tok = m.group(0)
        if tok.startswith("**"):
            parts.append((tok[2:-2], True, False))
        else:
            parts.append((tok[1:-1], False, True))
        pos = m.end()
    if pos < len(s):
        parts.append((s[pos:], False, False))
    return parts


def load_refs(path):
    """Read the [n]-prefixed reference lines, swap 22/23 so the list follows
    citation order ([22] Goedhart cyan, [23] Hirano StayGold)."""
    refs = {}
    with open(path) as f:
        for line in f:
            m = re.match(r"^\[(\d+)\]\s", line)
            if m:
                refs[int(m.group(1))] = line.strip()
    hirano, goedhart = refs[22], refs[23]          # current 22=Hirano, 23=Goedhart
    refs[22] = goedhart.replace("[23]", "[22]", 1)
    refs[23] = hirano.replace("[22]", "[23]", 1)
    # drop uncited tail references 29-34 (matplotlib, Yang, Banerjee, Cranfill,
    # Stepanenko, Romei); 1-28 are all cited so no renumbering is needed
    return [refs[n] for n in sorted(refs) if n <= 28]


def build_caption(dst, segments):
    """Return a fresh Normal-styled w:p element built from (text, bold) runs."""
    new_p = OxmlElement("w:p")
    para = Paragraph(new_p, dst)
    para.style = dst.styles["Normal"]
    for text, bold in segments:
        run = para.add_run(text)
        if bold:
            run.bold = True
    return new_p


def rebuild_ref_paragraph(para, ref_lines):
    """Replace all runs of `para` with clean runs built from ref_lines."""
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    joined = " ".join(ref_lines)
    for text, bold, ital in parse_inline(joined):
        run = para.add_run(text)
        if bold:
            run.bold = True
        if ital:
            run.italic = True


def replace_in_runs(para, old, new):
    """Replace first occurrence of `old` within a single run of `para`."""
    for r in para.runs:
        if old in r.text:
            r.text = r.text.replace(old, new, 1)
            return True
    return False


def main():
    shutil.copy(SRC_FRONT, OUT)
    dst = docx.Document(OUT)
    src = docx.Document(SRC_TAIL)
    ref_lines = load_refs(REF_SRC)

    # ---- front fixes on dst ----
    fixes = [
        ("Pieri et al. [31]", "[31]", "[11]"),                 # Pieri ref typo
        (", n = 56 crystal structures; Figure S1)", "Figure S1", "Figure 3"),
        ("provide useful tests of this interpretation (Table 2)", "(Table 2)", "(Table 1)"),
        ("Table 2. Three pair comparisons", "Table 2.", "Table 1."),
        ("class hierarchy summarized in Table 1.", "Table 1.", "Table 2."),
        ("Table 1. Per-unique-FP", "Table 1.", "Table 2."),
    ]
    applied = {f[0]: False for f in fixes}
    for p in dst.paragraphs:
        for marker, old, new in fixes:
            if not applied[marker] and marker in p.text:
                ok = replace_in_runs(p, old, new)
                applied[marker] = ok
                if not ok:
                    print("WARN front fix failed:", marker)
    for marker, done in applied.items():
        if not done:
            print("WARN front fix NOT located:", marker)

    sectPr = dst.element.body.find(qn("w:sectPr"))

    # ---- restore the three paragraphs the master dropped from the
    #      resting-position section (AausFP2, physical interpretation,
    #      B-factor/polar contacts); they re-cite refs [20] and [21] ----
    RESTORE_STARTS = (
        "AausFP2 (PDB 6S68)",
        "The physical interpretation is straightforward",
        "Two additional structural properties carry independent brightness",
    )
    restore = {s: None for s in RESTORE_STARTS}
    for child in src.element.body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        txt = Paragraph(child, src).text.strip()
        for s in RESTORE_STARTS:
            if restore[s] is None and txt.startswith(s):
                restore[s] = copy.deepcopy(child)
    # anchor: the "Cage size reports chromophore chemistry" heading in dst
    cage_anchor = None
    for p in dst.paragraphs:
        if p.style.name == "Heading 3" and p.text.strip().startswith(
            "Cage size reports chromophore chemistry"
        ):
            cage_anchor = p._element
            break
    assert cage_anchor is not None, "cage-size heading not found in master"
    for s in RESTORE_STARTS:
        assert restore[s] is not None, f"restore paragraph not found: {s}"
        cage_anchor.addprevious(restore[s])

    # ---- walk tail and copy from cyan paragraph onward ----
    OMIT_CAPTIONS = (
        "Figure 2. Sterically allowed",
        "Figure 5. The GFP",
        "Figure 6. Cage size varies",
    )
    DUP_S6 = "Naïve sign-test p-values are reported for reference only; all inferential"

    collecting = False
    n_paras = n_tables = 0
    for child in list(src.element.body.iterchildren()):
        if child.tag == qn("w:p"):
            para = Paragraph(child, src)
            txt = para.text.strip()
            if not collecting:
                if txt.startswith("Cyan FPs are the one exception"):
                    collecting = True
                else:
                    continue
            if not txt:
                continue
            if any(txt.startswith(c) for c in OMIT_CAPTIONS):
                continue
            if txt.startswith(DUP_S6):
                continue  # drop duplicated S6 sentence
            # place the Figure S3 caption at the end of Section S1
            if txt.startswith("S2. Stokes shift"):
                sectPr.addprevious(build_caption(dst, FIGS3_CAPTION))
                n_paras += 1
            new_p = copy.deepcopy(child)
            np = Paragraph(new_p, dst)
            # transforms
            if txt.startswith("Figure 3. Gatekeeper"):
                replace_in_runs(np, "Figure 3.", "Figure 2.")
            elif txt.startswith("Figure 4. Chromophore resting"):
                replace_in_runs(np, "Figure 4.", "Figure 3.")
            elif "I146F mutation reported by Goedhart et al. [23]" in txt:
                replace_in_runs(np, "[23]", "[22]")
            elif "rigid dimer interface that damps chromophore motion thermally [22]" in txt:
                replace_in_runs(np, "[22]", "[23]")
            elif "negative hula-twist diagonal of slope -0.85" in txt:
                np.add_run(FIGS3_REF)
            elif txt.startswith("[1] Tsien"):
                rebuild_ref_paragraph(np, ref_lines)
            sectPr.addprevious(new_p)
            n_paras += 1
        elif child.tag == qn("w:tbl"):
            if not collecting:
                continue
            new_t = copy.deepcopy(child)
            sectPr.addprevious(new_t)
            n_tables += 1

    dst.save(OUT)
    print(f"Wrote {OUT}: appended {n_paras} paragraphs, {n_tables} tables.")


if __name__ == "__main__":
    main()
