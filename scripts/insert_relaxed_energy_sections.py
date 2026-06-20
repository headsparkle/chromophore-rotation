#!/usr/bin/env python3
"""
insert_relaxed_energy_sections.py
=================================

Insert the relaxed-cage-flexibility and position-203 interaction-energy
additions into the manuscript master docx, following the established
edit pattern (backup first, match existing Heading 3 / Normal styles,
then run verify_manuscript_numbers.py).

Adds:
  (1) Results / gatekeeper: a paragraph quantifying the Thr203 -> Tyr203
      interaction as an energy (after the engineering-history paragraph).
  (2) Results / Validation: a paragraph reporting the relaxed-cage test
      (after the "purely steric and static" limitations paragraph).
  (3) Methods 2.5: two paragraphs describing the relaxed scan and the
      energy decomposition (after the forward-deletion-test paragraph).
  (4) SI section S11 + Figure S4 (relaxed energy surfaces, 4 panels).
  (5) SI section S12 + Figure S5 (position-203 interaction energy).

Idempotency guard: aborts if an S11 heading already exists.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
MASTER = ROOT / "manuscript-complete-mz-new.docx"
BACKUP = ROOT / "manuscript-complete-mz_backup-relaxenergy.docx"
FIG_S4 = ROOT / "figures" / "relaxed_dE_panel_S4.png"
FIG_S5 = ROOT / "figures" / "gatekeeper_energy_203.png"

GATEKEEPER_ANCHOR = "The engineering history of GFP also supports position 203."
VALIDATION_ANCHOR = "The scan is purely steric and static."
METHODS_ANCHOR = "Ile167Ala, served as a negative control."

P1_GATEKEEPER = (
    "We quantified the position-203 contact as a single-point non-bonded "
    "interaction energy between the residue-203 side chain and the chromophore "
    "phenol ring. In the 50 real Tyr203 yellow FPs the tyrosine ring sits in a "
    "near-parallel stack over the chromophore phenol, at a 3.6 Å ring-centroid "
    "distance and a 9° interplanar angle, with a van der Waals interaction near "
    "-3.2 kcal/mol. The Thr203 contact is weaker, near -0.8 kcal/mol, and carries "
    "no aromatic system. The T203Y substitution therefore converts a weak steric "
    "contact into a dispersion-bonded ground-state clamp, the energetic counterpart "
    "of the volume change measured by the deletion test. An in-silico T203Y graft "
    "onto the green Thr203 scaffolds deepens this van der Waals interaction in 92% "
    "of cases, and where the grafted ring relaxes into the observed stacking "
    "geometry it reproduces the interaction energy of the real yellow proteins. The "
    "electrostatic term is smaller and depends on the assumed chromophore "
    "protonation. Supplementary Section S12 and Figure S5 give the details."
)

P2_VALIDATION = (
    "One of these limits can be tested directly. We repeated the scan on four "
    "representative FPs (avGFP, mScarlet, mCherry, mTurquoise) with the chromophore "
    "held at each (τ, φ) while the side chains within 5 Å were allowed "
    "to relax under a soft-core potential. Letting the cage breathe widens the "
    "accessible region only slightly and never opens the I-bond wall. The energy "
    "minimum of the relaxed surface coincides with the deposited geometry, the side "
    "chains move by less than a few tenths of an Angstrom, and the bright/dim "
    "ordering of the four proteins is preserved. The result is insensitive to the "
    "two parameters of the relaxation. Supplementary Section S11 and Figure S4 give "
    "the relaxed energy surfaces."
)

P3A_METHODS = (
    "To test the rigid-cage assumption, a relaxed variant of the scan was run on "
    "four representative structures: avGFP (1EMA), mScarlet (5LK4), mCherry (3KCS), "
    "and mTurquoise (2YE0). At each (τ, φ) on a 10° grid the chromophore was "
    "held fixed and the side-chain heavy atoms of standard residues within 5 Å "
    "were energy-minimized under a soft-core Lennard-Jones potential, using the same "
    "Bondi radii as the rigid scan, together with a harmonic tether to the crystal "
    "position that stands in for the bonded network. Because the fused chromophore "
    "is held rigid, only its non-bonded interactions with the cage contribute, so "
    "this is a restrained soft-sphere relaxation rather than a full molecular-"
    "dynamics minimization. The Lennard-Jones well depth and the tether stiffness "
    "were each varied over a three-value grid to confirm that the conclusions do not "
    "depend on their values (Section S11)."
)

P3B_METHODS = (
    "The position-203 contact was also evaluated energetically. For each green "
    "Thr203 scaffold, the non-bonded interaction energy between the residue-203 side "
    "chain and the chromophore phenol ring was computed and separated into a van der "
    "Waals (Lennard-Jones) term and a Coulomb term, using AMBER parm10 van der Waals "
    "parameters and ff14SB partial charges. A Tyr203 side chain was then grafted "
    "onto the same backbone by superposing a deposited 1YFP rotamer, and the energy "
    "recomputed, so the within-scaffold difference isolates the T203Y change while "
    "leaving the chromophore untouched. The result was validated against the 50 real "
    "Tyr203 yellow FPs evaluated as deposited. The chromophore phenol was treated "
    "with both a neutral and an anionic-phenolate charge model to bracket the "
    "electrostatic term (Section S12)."
)

S11_HEAD = "S11. Robustness to cage flexibility"
S11_BODY = (
    "The rotational scan scores the chromophore against a rigid cage. To check that "
    "a rigid wall does not overstate how forbidden a rotamer is, we relaxed the cage "
    "on four representative FPs spanning the color range (Figure S4). At each "
    "(τ, φ) the chromophore was held fixed and the surrounding side chains within "
    "5 Å were minimized under a soft-core Lennard-Jones potential with a harmonic "
    "tether to the crystal position (Methods). Three features of the rigid analysis "
    "survive the relaxation. First, the energy minimum of each relaxed surface "
    "coincides with the deposited geometry, so the flexible-cage model independently "
    "recovers where each chromophore sits. Second, allowing the cage to breathe "
    "widens the accessible region only slightly and never opens the high-τ I-bond "
    "wall; the side chains move by less than a few tenths of an Angstrom, so this is "
    "a local relaxation rather than a rearrangement. Third, the bright/dim ordering "
    "is preserved: the dim, twisted red mCherry has by far the roomiest cage, with a "
    "low-energy basin centered on a twisted I-bond, while the bright FPs are confined "
    "to a narrow near-planar valley. A three-by-three sensitivity grid over the "
    "Lennard-Jones well depth (0.05 to 0.20 kcal/mol) and the tether stiffness (2 to "
    "20 kcal/mol per square Angstrom) leaves all three features unchanged: mCherry is "
    "the roomiest cage in every one of the nine settings, the relaxed minimum stays "
    "within a few kcal/mol of the deposited geometry, and the side-chain motion stays "
    "sub-Angstrom. The absolute accessible area scales smoothly with these "
    "parameters, so the relaxed scan supports the rank and shape conclusions of the "
    "rigid scan rather than a particular numerical area."
)
S11_CAP = (
    "Figure S4. Relaxed (cage-breathing) energy surfaces for four representative FPs. "
    "At each (τ, φ) on a 10° grid the chromophore is held fixed and the side chains "
    "within 5 Å are minimized under a soft-core Lennard-Jones potential. Color is the "
    "relaxed steric energy ΔE relative to the global minimum of each surface (capped "
    "at 25 kcal/mol). The white contour is the rigid-allowed boundary and the cyan "
    "dashed contour the relaxed-allowed boundary; the red star is the deposited "
    "(τ, φ). The dim, twisted red mCherry (3KCS) has the largest low-energy region, "
    "centered on a twisted I-bond, while the bright FPs are confined to a narrow "
    "near-planar valley."
)

S12_HEAD = "S12. Energetics of the position-203 / phenol interaction"
S12_BODY = (
    "The forward deletion test (main text) measures the volume that the position-203 "
    "side chain clears. To express the same effect as an energy, we computed the "
    "non-bonded interaction between the residue-203 side chain and the chromophore "
    "phenol ring and split it into van der Waals and Coulomb terms (Methods). A "
    "parallel aromatic stack is stabilized through Lennard-Jones dispersion and the "
    "electrostatic quadrupole, so we report the two terms and separately characterize "
    "the ring-ring geometry rather than treating π-stacking as a distinct term. In "
    "the 50 real Tyr203 yellow FPs the tyrosine ring stacks over the chromophore "
    "phenol at a median ring-centroid distance of 3.65 Å and an interplanar angle of "
    "9°, with a van der Waals interaction near -3.2 kcal/mol. The deposited Thr203 "
    "contact is near -0.8 kcal/mol and has no aromatic system. The within-scaffold "
    "in-silico T203Y deepens the van der Waals interaction in 92% of the 188 green "
    "Thr203 scaffolds; because the grafted rotamer is not relaxed it lands on average "
    "about 0.5 Å farther than the observed stacks and so gives a lower bound, but the "
    "30% of grafts that land in the observed stacking geometry reproduce the "
    "-3.2 kcal/mol of the real yellow proteins. The electrostatic term is model-"
    "dependent: with the bright-state anionic phenolate it adds roughly -3 kcal/mol "
    "to the T203Y change, while with a neutral phenol ring it is negligible. The van "
    "der Waals dispersion therefore carries the robust signal, consistent with the "
    "engineering of a dispersion-bonded ground-state clamp at position 203."
)
S12_CAP = (
    "Figure S5. Energetics of the position-203 / chromophore-phenol interaction. "
    "Left: distribution of the within-scaffold T203Y change in van der Waals "
    "interaction energy across 188 green Thr203 scaffolds (median -1.43 kcal/mol; "
    "92% favorable; six unrelaxed-graft clash outliers above +8 kcal/mol are off-"
    "scale). Right: Tyr203 / chromophore ring-stacking geometry, ring-centroid "
    "distance versus interplanar angle, colored by the Tyr203 van der Waals energy, "
    "with the 50 real Tyr203 yellow FPs overlaid as red circles. The real yellows "
    "cluster at the tightest, most parallel stacks (3.6 to 3.8 Å, under 15°)."
)


def find_para(doc, text):
    for p in doc.paragraphs:
        if text in p.text:
            return p
    raise SystemExit(f"anchor not found: {text!r}")


def new_para_after(ref: Paragraph, text="", style=None, image=None,
                   img_width_in=6.5):
    new_p = OxmlElement("w:p")
    ref._p.addnext(new_p)
    para = Paragraph(new_p, ref._parent)
    if style:
        para.style = style
    if image is not None:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(image), width=Inches(img_width_in))
    elif text:
        para.add_run(text)
    return para


def main():
    if not MASTER.is_file():
        raise SystemExit(f"master not found: {MASTER}")
    doc = Document(str(MASTER))
    if any(p.text.strip().startswith("S11.") for p in doc.paragraphs):
        raise SystemExit("S11 already present; aborting to avoid duplicate insert.")
    for f in (FIG_S4, FIG_S5):
        if not f.is_file():
            raise SystemExit(f"figure missing: {f}")

    shutil.copy2(MASTER, BACKUP)
    print(f"backup -> {BACKUP.name}")

    # (1) Results / gatekeeper
    gk = find_para(doc, GATEKEEPER_ANCHOR)
    new_para_after(gk, P1_GATEKEEPER, style="Normal")
    print("inserted (1) gatekeeper energy paragraph")

    # (2) Results / Validation
    val = find_para(doc, VALIDATION_ANCHOR)
    new_para_after(val, P2_VALIDATION, style="Normal")
    print("inserted (2) validation relaxed-cage paragraph")

    # (3) Methods 2.5 (two paragraphs, in order)
    m = find_para(doc, METHODS_ANCHOR)
    m = new_para_after(m, P3A_METHODS, style="Normal")
    new_para_after(m, P3B_METHODS, style="Normal")
    print("inserted (3) two methods paragraphs")

    # (4)+(5) SI sections before References: anchor on last paragraph before it
    refs = None
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name == "Heading 2" and p.text.strip() == "References":
            refs = doc.paragraphs[i - 1]
            break
    if refs is None:
        raise SystemExit("References heading not found")

    cur = refs
    cur = new_para_after(cur, S11_HEAD, style="Heading 3")
    cur = new_para_after(cur, S11_BODY, style="Normal")
    cur = new_para_after(cur, image=FIG_S4)
    cur = new_para_after(cur, S11_CAP, style="Normal")
    cur = new_para_after(cur, S12_HEAD, style="Heading 3")
    cur = new_para_after(cur, S12_BODY, style="Normal")
    cur = new_para_after(cur, image=FIG_S5)
    cur = new_para_after(cur, S12_CAP, style="Normal")
    print("inserted (4) S11 + Figure S4 and (5) S12 + Figure S5")

    doc.save(str(MASTER))
    print(f"saved -> {MASTER.name}")


if __name__ == "__main__":
    sys.exit(main())
