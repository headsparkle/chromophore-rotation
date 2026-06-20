#!/usr/bin/env python3
"""Render manuscript-complete.docx to markdown (manuscript.md) so the two stay
in lockstep. Preserves run-level bold/italic, converts tables to markdown,
turns the embedded-figure labels into placeholders, and splits the reference
paragraph into one entry per line. Run from the project root.
"""
import re
import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

SRC = "manuscript-complete-mz-new.docx"
OUT = "manuscript.md"

HEAD = {"Heading 1": "# ", "Heading 2": "## ", "Heading 3": "### "}


def runs_to_md(para):
    out = []
    for r in para.runs:
        t = r.text
        if not t:
            continue
        if t.strip() == "":
            out.append(t)
            continue
        lead = len(t) - len(t.lstrip())
        trail = len(t) - len(t.rstrip())
        core = t.strip()
        if r.bold:
            core = f"**{core}**"
        if r.italic:
            core = f"*{core}*"
        out.append(t[:lead] + core + (t[len(t) - trail:] if trail else ""))
    return "".join(out)


def _is_sep_row(cells):
    """True for a literal markdown separator row (e.g. '----') stored as data."""
    return all(c.strip() and set(c.strip()) <= set("-:") for c in cells)


def table_to_md(tbl):
    rows = tbl.rows
    lines = []
    header = [c.text.strip().replace("\n", " ") for c in rows[0].cells]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rows[1:]:
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        if _is_sep_row(cells):
            continue
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def main():
    d = docx.Document(SRC)
    paras = d.paragraphs
    tbls = d.tables
    pi = ti = 0
    blocks = []
    fig_label = re.compile(r"^Figure\s*(\d+):\s*(.*)$")
    for child in d.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = paras[pi]
            pi += 1
            txt = p.text.strip()
            if not txt:
                continue
            style = p.style.name
            if style in HEAD:
                blocks.append(HEAD[style] + txt)
                continue
            m = fig_label.match(txt)
            if m and len(txt) < 40:  # the embedded-figure placeholder labels
                num, desc = m.group(1), m.group(2).strip()
                label = f"Figure {num}" + (f": {desc}" if desc else "")
                blocks.append(
                    f"*[{label} (figure embedded in the DOCX; full caption in Figure Captions below).]*"
                )
                continue
            md = runs_to_md(p)
            if txt.startswith("[1] Tsien"):
                # split references one-per-line
                md = re.sub(r"\s*(\[\d+\])\s*", r"\n\1 ", md).strip()
            blocks.append(md)
        elif child.tag == qn("w:tbl"):
            tbl = tbls[ti]
            ti += 1
            blocks.append(table_to_md(tbl))

    text = "\n\n".join(blocks) + "\n"
    with open(OUT, "w") as fh:
        fh.write(text)
    print(f"Wrote {OUT}: {len(blocks)} blocks, {pi} paragraphs, {ti} tables.")


if __name__ == "__main__":
    main()
