#!/usr/bin/env python3
"""Edits on manuscript-complete-mz.docx (master):
1. Reference [12] -> Begg et al. ChemRxiv 2026.
2. Supplementary figures numbered by appearance: S1 = phi-tau (S1), S2 = Stokes
   (S2), S3 = island orientation (S9). Adds the S2 in-text mention, renumbers the
   S9 reference S2->S3, inserts S2 and S3 captions.
3. Ang -> Å, deg/degrees -> ° in prose and tables (protecting A_allowed_deg2 and
   "degree of freedom").
"""
import re, shutil, docx
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

MASTER = "manuscript-complete-mz.docx"
shutil.copy(MASTER, "manuscript-complete-mz_backup-supp.docx")
d = docx.Document(MASTER)

# ---------- 1. Reference [12] ----------
REF12 = [
    ("[12] Begg, L.P.; Mason, M.L.; Zimmer, M. Barrel Shape and Chromophore "
     "Rigidity Predict Fluorescent-Protein Photophysics. ", False, False),
    ("ChemRxiv", True, False),
    (" ", False, False),
    ("2026", False, True),
    (". DOI: 10.26434/chemrxiv.15003310/v2.", False, False),
]
for p in d.paragraphs:
    if p.text.strip().startswith("[12]"):
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        for t, i, b in REF12:
            run = p.add_run(t)
            run.italic, run.bold = i, b
        break

# ---------- 2. supplementary figure numbering ----------
def edit_run(pred, old, new):
    for p in d.paragraphs:
        if pred(p):
            for r in p.runs:
                if old in r.text:
                    r.text = r.text.replace(old, new, 1)
                    return True
    return False

# S9 island-orientation reference: Figure S2 -> Figure S3
assert edit_run(lambda p: "Island orientation." in p.text, "Figure S2", "Figure S3")
# S2 Stokes section: add an in-text mention
assert edit_run(
    lambda p: p.text.strip().startswith("The pooled Spearman correlation between Stokes shift"),
    "n = 644 structures)", "n = 644 structures; Figure S2)")

def find(pred):
    return next(p for p in d.paragraphs if pred(p))

def insert_caption(segs, anchor):
    newp = OxmlElement("w:p")
    anchor.addprevious(newp)
    para = Paragraph(newp, d)
    para.style = d.styles["Normal"]
    for text, bold in segs:
        run = para.add_run(text)
        run.bold = bold

FIGS2 = [
    ("Figure S2. Cage size versus Stokes shift.", True),
    (" Folded sterically allowed fraction f_allowed (folded) versus Stokes shift "
     "(nm), colored by color class. (a) All color classes (n = 644 structures): a "
     "weak negative pooled correlation (Spearman rho = -0.11, p = 0.006) that "
     "survives per-unique-FP aggregation (rho = -0.16, p = 0.04, n = 164); "
     "tight-caged red and blue FPs tend to have large Stokes shifts. (b) Green FPs "
     "only (n = 416 per PDB): an apparent within-green reversal (rho = +0.14, "
     "p = 0.004) that does not survive per-unique-FP aggregation (rho = +0.14, "
     "p = 0.26, n = 66) and is attributed to crystal-form oversampling.", False),
]
FIGS3 = [
    ("Figure S3. The sterically allowed (τ, φ) island is oriented almost "
     "exclusively along the P-bond axis.", True),
    (" For each of the 792 structures with at least three allowed cells under the "
     "soft (0.65 Å) clash tolerance, an ellipse was fit to the allowed-cell "
     "centroids and the angle of its major axis recorded (0° = τ axis; ±90° = φ "
     "axis; ±45° = hula-twist diagonals). (a) Histogram of major-axis angles: the "
     "distribution is sharply unimodal at +89° (IQR +88° to +89°), with 94% of "
     "structures within 15° of the pure φ direction, only 0.4% near the τ axis, "
     "and 3.3% near a hula-twist diagonal. (b) Polar view of the same angles. The "
     "allowed region elongates along the φ (P-bond) axis, confirming that the "
     "barrel leaves the P-bond as the variable rotational degree of freedom while "
     "clamping the I-bond.", False),
]
insert_caption(FIGS2, find(lambda p: p.style.name.startswith("Heading")
                           and p.text.strip().startswith("S3. Pseudoreplication"))._element)
insert_caption(FIGS3, find(lambda p: p.text.strip().startswith("Quantities reported per structure"))._element)

# ---------- 3. symbol fixes ----------
def fix(text):
    text = re.sub(r"\bAng\b", "Å", text)
    text = text.replace("deg^2", "°²")
    text = text.replace("(-90, 90] degrees", "(−90°, 90°]")
    text = text.replace("per degree", "per °")
    text = re.sub(r"(\d(?:[\d.]*\d)?)[\s\-]*degrees?\b", r"\1°", text)
    text = re.sub(r"(\d(?:[\d.]*\d)?)\s*deg\b", r"\1°", text)
    return text

def fix_runs(p):
    for r in p.runs:
        nt = fix(r.text)
        if nt != r.text:
            r.text = nt

for p in d.paragraphs:
    fix_runs(p)
for tbl in d.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                fix_runs(p)

d.save(MASTER)
print("done: ref[12], supp figures S1/S2/S3, symbol conversions")
