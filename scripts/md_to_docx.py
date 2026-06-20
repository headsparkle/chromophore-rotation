"""
md_to_docx.py -- convert manuscript.md to manuscript.docx using python-docx.
Handles: headings, bold, italic, inline code, bullet/numbered lists, tables,
         horizontal rules, and block paragraphs.
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT = Path(__file__).resolve().parent.parent
MD_FILE  = PROJECT / "manuscript.md"
OUT_FILE = PROJECT / "manuscript.docx"


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_runs(para, text):
    """Parse inline markdown (bold, italic, inline code) and add runs."""
    # Pattern: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        # plain text before match
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        if m.group(0).startswith('**'):
            run = para.add_run(m.group(2))
            run.bold = True
        elif m.group(0).startswith('*'):
            run = para.add_run(m.group(3))
            run.italic = True
        elif m.group(0).startswith('`'):
            run = para.add_run(m.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


def parse_table(doc, table_lines):
    """Render a markdown table into a docx table."""
    rows = []
    for line in table_lines:
        if re.match(r'^\s*\|[\s\-:]+\|\s*$', line):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= ncols:
                break
            cell = tbl.cell(i, j)
            cell.text = ''
            para = cell.paragraphs[0]
            add_runs(para, cell_text)
            if i == 0:
                for run in para.runs:
                    run.bold = True


def md_to_docx(md_path, out_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    lines = md_path.read_text(encoding='utf-8').splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # ---- Heading
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            heading_map = {1: 'Heading 1', 2: 'Heading 2',
                           3: 'Heading 3', 4: 'Heading 4'}
            para = doc.add_paragraph(style=heading_map.get(level, 'Heading 3'))
            add_runs(para, text)
            i += 1
            continue

        # ---- Horizontal rule
        if re.match(r'^[-*]{3,}\s*$', line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ---- Table: collect consecutive table lines
        if re.match(r'^\s*\|', line):
            table_lines = []
            while i < len(lines) and re.match(r'^\s*\|', lines[i]):
                table_lines.append(lines[i])
                i += 1
            parse_table(doc, table_lines)
            continue

        # ---- Blank line
        if not line.strip():
            i += 1
            continue

        # ---- Numbered list item
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            para = doc.add_paragraph(style='List Number')
            add_runs(para, m.group(2))
            i += 1
            continue

        # ---- Bullet list item (-, *, or indented -)
        m = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m:
            indent = len(m.group(1)) // 2
            style_name = 'List Bullet' if indent == 0 else 'List Bullet 2'
            para = doc.add_paragraph(style=style_name)
            add_runs(para, m.group(2))
            i += 1
            continue

        # ---- Regular paragraph (may span multiple non-blank lines)
        para_lines = []
        while i < len(lines) and lines[i].strip() \
              and not re.match(r'^#', lines[i]) \
              and not re.match(r'^\s*\|', lines[i]) \
              and not re.match(r'^[-*]{3,}\s*$', lines[i]) \
              and not re.match(r'^(\d+)\.\s', lines[i]) \
              and not re.match(r'^(\s*)[-*]\s', lines[i]):
            para_lines.append(lines[i].strip())
            i += 1
        if para_lines:
            para = doc.add_paragraph(style='Normal')
            add_runs(para, ' '.join(para_lines))

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    md_to_docx(MD_FILE, OUT_FILE)
