"""
insert_eval_confirmatory.py
===========================
Additively insert the confirmatory analyses requested by the external evaluation
into manuscript-complete-mz.docx. All additions are PROSE with new numbers; the
headline numbers checked by verify_manuscript_numbers.py are untouched.

Blocks:
  (A) main-text Validation paragraph: append one robustness-summary sentence
      surfacing permutation p + Theil-Sen + jackknife/outlier grid.
  (B) S7: new "Resting twist: P-bond versus I-bond." paragraph after the cage
      "P/I decomposition." paragraph.
  (C) S8: new "Permutation inference, robust slope, and outlier sensitivity."
      paragraph after the AausFP2 note that closes the S8 robustness table.
  (D) S8: new "Heteroatoms and ions." paragraph (heteroatom-exclusion rescan).

Numbers from data/eval_confirmatory.csv and data/heteroatom_sensitivity.csv.
Backs up first. No em-dashes.
"""
from __future__ import annotations
import copy
import shutil
from pathlib import Path

import docx

PROJECT = Path(__file__).resolve().parent.parent
SRC = PROJECT / "manuscript-complete-mz.docx"
BACKUP = PROJECT / "manuscript-complete-mz_backup-evalconf.docx"

# ---- anchors (substring match on paragraph text) -------------------------
A_ANCHOR = "Full results are in Supplementary Sections S5 and S6"
B_ANCHOR = "P/I decomposition. Separating f_allowed"
# S8 AausFP2 note (unique to S8; the Fig.4 caption also names AausFP2/6S68,
# so anchor on the S8-only phrase "does not appear in any row")
C_ANCHOR = "the labeled endpoint in Figure 4, does not appear in any row"
D_ANCHOR = C_ANCHOR  # D inserted right after C

# ---- (A) main-text appended sentence -------------------------------------
A_SENTENCE = (
    " The within-red twist association is robust to the choices that most "
    "concern small-sample inference: a 20,000-shuffle permutation test "
    "reproduces the asymptotic significance (permutation p = 0.0002 per crystal, "
    "p = 0.035 per unique FP), the Theil-Sen robust slope is negative with a 95% "
    "confidence interval entirely below zero in both the per-crystal and "
    "per-unique views, and a leave-one-protein-out jackknife holds the rank "
    "correlation between -0.41 and -0.30. The trend is not the work of a few "
    "well-known proteins, although the bright, near-planar mScarlet group does "
    "contribute genuine support (Supplementary Section S8)."
)

# ---- (B) S7 resting-twist decomposition ----------------------------------
B_BOLD = "Resting twist: P-bond versus I-bond."
B_BODY = (
    " The cage decomposition above asks which bond the barrel leaves free; a "
    "separate question is which bond's resting deviation carries the red "
    "brightness signal. Splitting d_exp_to_planar into its I-bond (absolute tau "
    "deviation) and P-bond (absolute phi deviation) components, the red "
    "twist/QY association is carried by the P-bond coordinate: per crystal the "
    "P-bond deviation correlates with QY at rho = -0.38 (permutation p = 0.004, "
    "n = 56) while the I-bond deviation does not (rho = -0.22, p = 0.11). This is "
    "the expected internal-consistency check for the I-bond-clamp model. Because "
    "the I-bond is already clamped near planar in nearly every barrel, it has "
    "little resting spread to correlate, so the variation that tracks brightness "
    "lies on the P-bond side. The decomposition has power only in the per-crystal "
    "view; per unique FP (n = 38) both components are individually weak (P-bond "
    "rho = -0.20, I-bond rho = -0.10), so only the combined distance d_exp_to_planar "
    "is reported as the primary metric."
)

# ---- (C) S8 permutation / robust / outlier -------------------------------
C_BOLD = "Permutation inference, robust slope, and outlier sensitivity."
C_BODY = (
    " Because the red subsets are on the order of tens of proteins, the headline "
    "association was confirmed with three small-sample-appropriate checks. First, "
    "a 20,000-shuffle permutation test gives essentially the same significance as "
    "the asymptotic Spearman p (red per crystal permutation p = 0.0002 versus "
    "asymptotic 0.0001; per unique FP 0.035 versus 0.034; NRQ 0.026; CRQ 0.019), "
    "so the asymptotic p-values are reliable for this dataset. Second, the "
    "Theil-Sen robust slope of log10(QY) on d_exp_to_planar is negative with a 95% "
    "confidence interval entirely below zero in both views (per crystal -0.014 "
    "[-0.023, -0.006]; per unique FP -0.011 [-0.024, -0.002]), so the trend is not "
    "an artifact of a few high-leverage points. Third, a leave-one-protein-out "
    "jackknife on the per-unique cohort holds the rank correlation within [-0.41, "
    "-0.30], and a named outlier-deletion grid shows that removing mRouge alone, "
    "or the single most twisted FP, leaves the result unchanged (rho = -0.345, "
    "p = 0.035 in each case), whereas removing the entire mScarlet group of five "
    "bright, near-planar FPs weakens it to rho = -0.28 (p = 0.12). The bright "
    "near-planar end therefore contributes real support, but the association does "
    "not depend on the twisted tail or on any single protein."
)

# ---- (D) heteroatom rescan -- FILLED IN AFTER RESCAN ---------------------
D_BOLD = "Heteroatoms and ions."
D_BODY = (
    " The production cage includes every non-water heavy atom, so bound ions, "
    "buffer molecules and crystallographic ligands are part of the steric cage. "
    "To test whether they affect the result, we rescanned all 93 red crystal "
    "structures with every non-polymer heteroatom removed from the cage (water "
    "and non-polymer het excluded; modified polymer residues such as "
    "selenomethionine retained). Of the 93, 43 carry at least one non-polymer "
    "heteroatom and 6 have one within 5 A of the chromophore (closest 2.8 A, in "
    "3PIB), yet f_allowed_folded is unchanged in every structure (maximum change "
    "below 0.0001), because the heteroatoms do not lie in the volume the rotating "
    "phenol sweeps. The within-red cage-size/QY null is identical under both cages "
    "(rho = +0.10 per crystal either way). Heteroatom inclusion therefore has no "
    "effect on any reported cage metric or conclusion."
)


def insert_para_after(anchor_para, bold_text, body_text):
    """Insert a new paragraph (bold lead-in run + normal body) after anchor."""
    new_p = copy.deepcopy(anchor_para._p)
    # clear copied runs
    for r in list(new_p.findall(docx.oxml.ns.qn("w:r"))):
        new_p.remove(r)
    anchor_para._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, anchor_para._parent)
    if bold_text:
        run = para.add_run(bold_text)
        run.bold = True
    para.add_run(body_text)
    return para


def main():
    have_D = D_BODY != "__FILL_AFTER_RESCAN__"
    shutil.copy(SRC, BACKUP)
    d = docx.Document(str(SRC))
    paras = d.paragraphs

    def find(anchor):
        for p in paras:
            if anchor in p.text:
                return p
        raise RuntimeError(f"anchor not found: {anchor!r}")

    done = []
    # A: append sentence to Validation paragraph (idempotent guard)
    a = find(A_ANCHOR)
    if "20,000-shuffle permutation test reproduces" not in a.text:
        a.add_run(A_SENTENCE)
        done.append("A")

    # B: after cage P/I decomposition
    if not any(B_BOLD in p.text for p in paras):
        insert_para_after(find(B_ANCHOR), B_BOLD, B_BODY)
        done.append("B")

    c_anchor = find(C_ANCHOR)
    # C goes immediately after the S8 AausFP2 note; D goes after C.
    if not any(C_BOLD in p.text for p in paras):
        c_para = insert_para_after(c_anchor, C_BOLD, C_BODY)
        done.append("C")
    else:
        c_para = next(p for p in d.paragraphs if C_BOLD in p.text)
    if have_D and not any(D_BOLD in p.text for p in d.paragraphs):
        insert_para_after(c_para, D_BOLD, D_BODY)
        done.append("D")

    d.save(str(SRC))
    print(f"inserted {done or '(nothing new)'}. backup: {BACKUP.name}")


if __name__ == "__main__":
    main()
